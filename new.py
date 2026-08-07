import os
import sys
import io
import time
import docx
import openpyxl
from openpyxl.drawing.image import Image as OpenpyxlImage
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, Border, Side, Font, PatternFill
from PIL import Image as PILImage
import tkinter as tk
from tkinter import filedialog

def select_word_file():
    """
    Membuka jendela pemilih file interaktif (File Explorer).
    """
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)

    print("[INFO] Membuka jendela pemilih berkas... Silakan pilih dokumen Word (.docx)")
    file_path = filedialog.askopenfilename(
        title="Pilih Dokumen Word Logbook (.docx) yang Ingin Dipindahkan ke Excel",
        filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
    )
    root.destroy()
    return file_path

def convert_word_tables_to_excel(doc_path):
    if not doc_path or not os.path.exists(doc_path):
        print("[INFO] Tidak ada file yang dipilih atau file tidak ditemukan. Proses dibatalkan.")
        return

    doc_name = os.path.splitext(os.path.basename(doc_path))[0]
    output_folder = os.path.dirname(doc_path) or os.getcwd()
    output_excel_path = os.path.join(output_folder, f"Hasil_Otomatis_{doc_name}.xlsx")

    print(f"\n[INFO] Membuka dokumen Word & merapikan rasio foto alami: {doc_path}")
    doc = docx.Document(doc_path)

    if len(doc.tables) == 0:
        print("[WARNING] Dokumen ini tidak memiliki tabel!")
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Logbook"

    # Pastikan Garis Kisi (Gridlines) Excel Selalu Tampil
    ws.views.sheetView[0].showGridLines = True

    # Styling Font & Border
    font_header = Font(name="Calibri", size=10, bold=True, color="1F2937")
    font_body = Font(name="Calibri", size=10, color="1F2937")
    fill_header = PatternFill(start_color="F3F4F6", end_color="F3F4F6", fill_type="solid")

    thin_border = Border(
        left=Side(style='thin', color='9CA3AF'),
        right=Side(style='thin', color='9CA3AF'),
        top=Side(style='thin', color='9CA3AF'),
        bottom=Side(style='thin', color='9CA3AF')
    )

    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)

    current_excel_row = 1

    for table_idx, table in enumerate(doc.tables, start=1):
        print(f"[INFO] Memproses Tabel Logbook #{table_idx} ({len(table.rows)} baris)...")

        num_cols = len(table.columns)
        col_max_widths = {c: 16 for c in range(1, num_cols + 1)}

        for r_idx, row in enumerate(table.rows):
            is_header = (r_idx == 0)
            max_row_height = 26 if is_header else 32 # Default tinggi baris awal

            for c_idx, cell in enumerate(row.cells, start=1):
                col_letter = get_column_letter(c_idx)
                cell_ref = f"{col_letter}{current_excel_row}"
                excel_cell = ws.cell(row=current_excel_row, column=c_idx)

                # 1. Tulis Teks Sel & Styling
                cell_text = cell.text.strip()
                excel_cell.value = cell_text
                excel_cell.border = thin_border
                excel_cell.alignment = align_center

                if is_header:
                    excel_cell.font = font_header
                    excel_cell.fill = fill_header
                else:
                    excel_cell.font = font_body

                if cell_text:
                    lines = cell_text.split('\n')
                    longest_line = max([len(l) for l in lines]) if lines else 0
                    col_max_widths[c_idx] = max(col_max_widths.get(c_idx, 16), min(longest_line + 4, 35))

                # 2. Parsing Gambar di Sel
                images_in_cell = []
                for p in cell.paragraphs:
                    for r in p.runs:
                        blips = r._r.xpath('.//a:blip')
                        for blip in blips:
                            rId = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if rId and rId in cell.part.related_parts:
                                image_part = cell.part.related_parts[rId]
                                images_in_cell.append(image_part.image.blob)

                # 3. Penataan Foto Proporsional (Asli Tanpa Pipih / Terdistorsi)
                if images_in_cell:
                    for img_bytes in images_in_cell:
                        try:
                            pil_img = PILImage.open(io.BytesIO(img_bytes))

                            # Murni menjaga Aspect Ratio foto (maksimum 140x110 px)
                            max_w, max_h = 140, 110
                            pil_img.thumbnail((max_w, max_h), PILImage.Resampling.LANCZOS)

                            img_buffer = io.BytesIO()
                            pil_img.save(img_buffer, format='PNG')
                            img_buffer.seek(0)

                            # Gunakan Penempatan Anchor Standar Sel (Menjaga Rasio Asli)
                            xl_img = OpenpyxlImage(img_buffer)
                            ws.add_image(xl_img, cell_ref)

                            # Sesuaikan tinggi baris & lebar kolom agar foto muat pas tanpa terpotong
                            img_height_pt = (pil_img.height * 0.75) + 16
                            max_row_height = max(max_row_height, img_height_pt)

                            needed_w = (pil_img.width / 7) + 4
                            col_max_widths[c_idx] = max(col_max_widths.get(c_idx, 16), needed_w)

                        except Exception as e:
                            print(f"[WARNING] Gagal memproses gambar pada sel {cell_ref}: {e}")

            ws.row_dimensions[current_excel_row].height = max_row_height
            current_excel_row += 1

        for c_idx, width in col_max_widths.items():
            col_letter = get_column_letter(c_idx)
            ws.column_dimensions[col_letter].width = width

        current_excel_row += 1

    # Penanganan Safe Save (Permission Error Safe)
    saved_path = output_excel_path
    try:
        wb.save(output_excel_path)
    except PermissionError:
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        saved_path = os.path.join(output_folder, f"Hasil_Otomatis_{doc_name}_{timestamp}.xlsx")
        wb.save(saved_path)
        print(f"\n[CATATAN] File '{os.path.basename(output_excel_path)}' sedang dibuka di Microsoft Excel.")
        print(f"[SOLUSI] Hasil baru otomatis disimpan sebagai nama alternatif.")

    print("\n" + "=" * 65)
    print(f" SELESAI! Tabel Logbook & Foto Rasio Almi Berhasil Disimpan:")
    print(f" 📄 Input  : {doc_path}")
    print(f" 📊 Output : {os.path.abspath(saved_path)}")
    print("=" * 65)

if __name__ == "__main__":
    if len(sys.argv) > 1 and os.path.exists(sys.argv[1]):
        selected_file = sys.argv[1]
    else:
        selected_file = select_word_file()

    convert_word_tables_to_excel(selected_file)
